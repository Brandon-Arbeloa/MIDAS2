"""
Document Processing Background Tasks
Handles asynchronous document processing for MIDAS on Windows
"""

import os
import sys
import time
import hashlib
import traceback
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any
import tempfile
import shutil

from celery import Task
from celery.utils.log import get_task_logger
from celery_config import app

# Import MIDAS modules
sys.path.append(str(Path(__file__).parent.parent))
from document_processor_refactored import DocumentProcessor
from document_indexer import DocumentIndexingSystem
from config import ConfigManager

# Windows-specific imports
if sys.platform == 'win32':
    import win32file
    import win32con
    import pywintypes

logger = get_task_logger(__name__)

class DocumentProcessingTask(Task):
    """Base task class with Windows-specific error handling"""
    
    autoretry_for = (
        OSError,
        IOError,
        ConnectionError,
        TimeoutError,
    )
    retry_kwargs = {'max_retries': 3, 'countdown': 60}
    retry_backoff = True
    retry_jitter = True
    
    def on_failure(self, exc, task_id, args, kwargs, einfo):
        """Handle task failure with Windows-specific logging"""
        logger.error(f"Task {task_id} failed: {exc}")
        logger.error(f"Traceback: {einfo}")
        
        # Windows-specific error codes
        if sys.platform == 'win32' and isinstance(exc, OSError):
            error_code = exc.winerror if hasattr(exc, 'winerror') else None
            if error_code:
                logger.error(f"Windows error code: {error_code}")
                
                # Common Windows error codes
                error_messages = {
                    5: "Access denied - check file permissions",
                    32: "File is being used by another process",
                    53: "Network path not found",
                    67: "Network name not found",
                    1223: "Operation cancelled by user"
                }
                
                if error_code in error_messages:
                    logger.error(f"Windows error: {error_messages[error_code]}")

@app.task(base=DocumentProcessingTask, bind=True, name='tasks.document.process_file')
def process_document_file(self, file_path: str, user_id: int = None, 
                         options: Dict = None) -> Dict[str, Any]:
    """
    Process a single document file asynchronously
    
    Args:
        file_path: Path to the document file
        user_id: Optional user ID for tracking
        options: Processing options
        
    Returns:
        Dict with processing results
    """
    start_time = time.time()
    result = {
        'status': 'pending',
        'file_path': file_path,
        'user_id': user_id,
        'task_id': self.request.id,
        'errors': [],
        'chunks_created': 0,
        'processing_time': 0
    }
    
    try:
        file_path = Path(file_path)
        
        # Validate file exists
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > 100 * 1024 * 1024:  # 100MB limit
            raise ValueError(f"File too large: {file_size} bytes")
        
        # Update task state
        self.update_state(
            state='PROCESSING',
            meta={
                'current': 0,
                'total': 100,
                'status': f'Processing {file_path.name}'
            }
        )
        
        # Initialize processors
        config_manager = ConfigManager()
        config = config_manager.load_config()
        doc_processor = DocumentProcessor(config)
        
        # Windows file locking check
        if sys.platform == 'win32':
            if not check_file_access(file_path):
                raise OSError("File is locked by another process")
        
        # Create temporary copy for processing
        with tempfile.NamedTemporaryFile(
            suffix=file_path.suffix,
            delete=False
        ) as tmp_file:
            shutil.copy2(file_path, tmp_file.name)
            tmp_path = Path(tmp_file.name)
        
        try:
            # Process document
            self.update_state(
                state='PROCESSING',
                meta={
                    'current': 30,
                    'total': 100,
                    'status': 'Extracting content'
                }
            )
            
            chunks = doc_processor.process_file(tmp_path)
            result['chunks_created'] = len(chunks)
            
            # Index chunks if indexing system available
            self.update_state(
                state='PROCESSING',
                meta={
                    'current': 60,
                    'total': 100,
                    'status': 'Indexing content'
                }
            )
            
            try:
                indexing_system = DocumentIndexingSystem()
                indexed_ids = indexing_system.index_document(
                    str(file_path),
                    chunks,
                    metadata={
                        'user_id': user_id,
                        'processed_at': datetime.now().isoformat(),
                        'task_id': self.request.id
                    }
                )
                result['indexed_chunks'] = len(indexed_ids)
            except Exception as e:
                logger.warning(f"Indexing failed: {e}")
                result['errors'].append(f"Indexing error: {str(e)}")
            
            # Calculate file hash
            file_hash = calculate_file_hash(file_path)
            result['file_hash'] = file_hash
            
            # Update final state
            self.update_state(
                state='PROCESSING',
                meta={
                    'current': 90,
                    'total': 100,
                    'status': 'Finalizing'
                }
            )
            
            result['status'] = 'completed'
            result['processing_time'] = time.time() - start_time
            
        finally:
            # Clean up temporary file
            try:
                tmp_path.unlink()
            except Exception:
                pass
        
        logger.info(f"Successfully processed {file_path}: {result['chunks_created']} chunks")
        return result
        
    except Exception as e:
        logger.error(f"Error processing {file_path}: {e}")
        result['status'] = 'failed'
        result['errors'].append(str(e))
        result['processing_time'] = time.time() - start_time
        
        # Re-raise for retry mechanism
        raise

@app.task(bind=True, name='tasks.document.process_batch')
def process_document_batch(self, file_paths: List[str], user_id: int = None,
                          options: Dict = None) -> Dict[str, Any]:
    """Process multiple documents in batch"""
    
    results = {
        'total': len(file_paths),
        'completed': 0,
        'failed': 0,
        'task_id': self.request.id,
        'results': []
    }
    
    for i, file_path in enumerate(file_paths):
        self.update_state(
            state='PROCESSING',
            meta={
                'current': i,
                'total': len(file_paths),
                'status': f'Processing file {i+1} of {len(file_paths)}'
            }
        )
        
        try:
            # Process each file as a subtask
            task_result = process_document_file.apply_async(
                args=[file_path, user_id, options],
                queue='documents'
            )
            
            # Wait for result with timeout
            file_result = task_result.get(timeout=300)
            
            results['results'].append(file_result)
            if file_result['status'] == 'completed':
                results['completed'] += 1
            else:
                results['failed'] += 1
                
        except Exception as e:
            logger.error(f"Batch processing error for {file_path}: {e}")
            results['failed'] += 1
            results['results'].append({
                'file_path': file_path,
                'status': 'failed',
                'error': str(e)
            })
    
    return results

@app.task(name='tasks.document.extract_metadata')
def extract_document_metadata(file_path: str) -> Dict[str, Any]:
    """Extract metadata from document without full processing"""
    
    file_path = Path(file_path)
    metadata = {
        'file_name': file_path.name,
        'file_size': file_path.stat().st_size,
        'file_type': file_path.suffix.lower(),
        'created_at': datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
        'modified_at': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
    }
    
    # Extract additional metadata based on file type
    if file_path.suffix.lower() == '.pdf':
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                metadata['page_count'] = len(reader.pages)
                if reader.metadata:
                    metadata['pdf_metadata'] = {
                        'title': reader.metadata.get('/Title', ''),
                        'author': reader.metadata.get('/Author', ''),
                        'subject': reader.metadata.get('/Subject', ''),
                    }
        except Exception as e:
            logger.warning(f"Failed to extract PDF metadata: {e}")
    
    elif file_path.suffix.lower() in ['.docx', '.doc']:
        try:
            from docx import Document
            doc = Document(file_path)
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
        except Exception as e:
            logger.warning(f"Failed to extract Word metadata: {e}")
    
    return metadata

@app.task(name='tasks.document.cleanup_old_files')
def cleanup_old_processed_files(days_old: int = 7) -> Dict[str, int]:
    """Clean up old processed files from temporary storage"""
    
    temp_dir = Path(tempfile.gettempdir()) / 'midas_processing'
    if not temp_dir.exists():
        return {'deleted_files': 0, 'freed_space': 0}
    
    deleted_count = 0
    freed_space = 0
    cutoff_date = datetime.now().timestamp() - (days_old * 24 * 60 * 60)
    
    for file_path in temp_dir.glob('**/*'):
        if file_path.is_file():
            if file_path.stat().st_mtime < cutoff_date:
                try:
                    file_size = file_path.stat().st_size
                    file_path.unlink()
                    deleted_count += 1
                    freed_space += file_size
                except Exception as e:
                    logger.warning(f"Failed to delete {file_path}: {e}")
    
    return {
        'deleted_files': deleted_count,
        'freed_space': freed_space,
        'freed_space_mb': freed_space / (1024 * 1024)
    }

# Utility functions

def check_file_access(file_path: Path) -> bool:
    """Check if file is accessible on Windows"""
    if sys.platform != 'win32':
        return True
    
    try:
        handle = win32file.CreateFile(
            str(file_path),
            win32con.GENERIC_READ,
            win32con.FILE_SHARE_READ | win32con.FILE_SHARE_WRITE,
            None,
            win32con.OPEN_EXISTING,
            0,
            None
        )
        win32file.CloseHandle(handle)
        return True
    except pywintypes.error:
        return False

def calculate_file_hash(file_path: Path) -> str:
    """Calculate SHA256 hash of file"""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()

# Export tasks
__all__ = [
    'process_document_file',
    'process_document_batch',
    'extract_document_metadata',
    'cleanup_old_processed_files'
]