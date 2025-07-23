"""
Document Indexing System for MIDAS RAG - Windows 11 Optimized
Processes CSV, JSON, TXT, PDF, and DOCX files with intelligent chunking
Uses sentence-transformers for local embeddings and Qdrant for vector storage
"""

import os
import json
import csv
import logging
import hashlib
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Generator
from datetime import datetime, timezone
import time
import gc
import traceback

# Core libraries
import pandas as pd
import numpy as np
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import VectorParams, Distance, PointStruct
from qdrant_client.http import models

# Document processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Windows-specific imports
import sys
if sys.platform == "win32":
    import win32file
    import win32con
    WINDOWS_SPECIFIC = True
else:
    WINDOWS_SPECIFIC = False


class WindowsLogger:
    """Windows-specific logger with file rotation and proper path handling"""
    
    def __init__(self, log_dir: str = "C:\\MIDAS\\logs", name: str = "document_indexer"):
        self.log_dir = Path(log_dir)
        self.log_dir.mkdir(parents=True, exist_ok=True)
        self.name = name
        
        # Setup logger
        self.logger = logging.getLogger(name)
        self.logger.setLevel(logging.INFO)
        
        # Clear existing handlers
        self.logger.handlers.clear()
        
        # File handler with Windows-safe filename
        timestamp = datetime.now().strftime("%Y%m%d")
        log_file = self.log_dir / f"{name}_{timestamp}.log"
        
        file_handler = logging.FileHandler(log_file, encoding='utf-8')
        file_handler.setLevel(logging.INFO)
        
        # Console handler
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        
        # Formatter
        formatter = logging.Formatter(
            '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        
        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)
    
    def info(self, message: str):
        self.logger.info(message)
    
    def error(self, message: str):
        self.logger.error(message)
    
    def warning(self, message: str):
        self.logger.warning(message)
    
    def debug(self, message: str):
        self.logger.debug(message)


class WindowsFileHandler:
    """Windows-specific file handling with proper path management"""
    
    @staticmethod
    def is_file_locked(file_path: Path) -> bool:
        """Check if a file is locked (Windows-specific)"""
        if not WINDOWS_SPECIFIC:
            return False
            
        try:
            handle = win32file.CreateFile(
                str(file_path),
                win32con.GENERIC_READ,
                0,  # No sharing
                None,
                win32con.OPEN_EXISTING,
                win32con.FILE_ATTRIBUTE_NORMAL,
                None
            )
            win32file.CloseHandle(handle)
            return False
        except:
            return True
    
    @staticmethod
    def get_file_metadata(file_path: Path) -> Dict[str, Any]:
        """Get comprehensive file metadata with Windows attributes"""
        try:
            stat = file_path.stat()
            metadata = {
                'file_path': str(file_path.resolve()),
                'file_name': file_path.name,
                'file_extension': file_path.suffix.lower(),
                'file_size': stat.st_size,
                'created_time': datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat(),
                'modified_time': datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
                'accessed_time': datetime.fromtimestamp(stat.st_atime, tz=timezone.utc).isoformat(),
                'mime_type': mimetypes.guess_type(str(file_path))[0] or 'application/octet-stream'
            }
            
            # Add Windows-specific attributes if available
            if WINDOWS_SPECIFIC:
                try:
                    import win32api
                    metadata['file_attributes'] = win32api.GetFileAttributes(str(file_path))
                except:
                    pass
            
            return metadata
        except Exception as e:
            return {
                'file_path': str(file_path),
                'error': str(e)
            }
    
    @staticmethod
    def safe_read_text_file(file_path: Path, encodings: List[str] = None) -> Optional[str]:
        """Safely read text file with multiple encoding attempts"""
        if encodings is None:
            encodings = ['utf-8', 'utf-8-sig', 'cp1252', 'iso-8859-1', 'latin1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                break
        
        return None


class DocumentChunker:
    """Intelligent document chunking with overlap for optimal retrieval"""
    
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.logger = WindowsLogger().logger
    
    def chunk_text(self, text: str, document_metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Chunk text intelligently with overlap and metadata preservation
        
        Args:
            text: The text content to chunk
            document_metadata: Metadata about the source document
            
        Returns:
            List of chunk dictionaries with text and metadata
        """
        if not text or not text.strip():
            return []
        
        chunks = []
        
        # Try sentence-based chunking first
        try:
            sentences = self._split_into_sentences(text)
            chunks = self._create_sentence_chunks(sentences, document_metadata)
        except Exception as e:
            self.logger.warning(f"Sentence chunking failed: {e}, falling back to word chunking")
            chunks = self._create_word_chunks(text, document_metadata)
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using simple heuristics"""
        import re
        
        # Clean up the text
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Split on sentence endings, but be careful with abbreviations
        sentence_endings = re.compile(r'([.!?]+)(\s+[A-Z])')
        sentences = []
        
        last_end = 0
        for match in sentence_endings.finditer(text):
            sentence = text[last_end:match.start() + len(match.group(1))].strip()
            if sentence:
                sentences.append(sentence)
            last_end = match.start() + len(match.group(1))
        
        # Add remaining text
        remaining = text[last_end:].strip()
        if remaining:
            sentences.append(remaining)
        
        return [s for s in sentences if len(s.strip()) > 10]  # Filter very short sentences
    
    def _create_sentence_chunks(self, sentences: List[str], document_metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Create chunks from sentences with intelligent overlap"""
        if not sentences:
            return []
        
        chunks = []
        current_chunk = []
        current_length = 0
        chunk_index = 0
        
        i = 0
        while i < len(sentences):
            sentence = sentences[i]
            sentence_length = len(sentence.split())
            
            # If adding this sentence would exceed chunk size and we have content
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Create chunk
                chunk_text = ' '.join(current_chunk)
                chunk_data = self._create_chunk_metadata(
                    chunk_text, chunk_index, document_metadata, 
                    len(current_chunk), current_length
                )
                chunks.append(chunk_data)
                
                # Calculate overlap
                overlap_sentences = self._calculate_sentence_overlap(current_chunk)
                current_chunk = overlap_sentences
                current_length = sum(len(s.split()) for s in current_chunk)
                chunk_index += 1
                
                # Don't increment i, try to add the same sentence to new chunk
                continue
            
            # Add sentence to current chunk
            current_chunk.append(sentence)
            current_length += sentence_length
            i += 1
        
        # Add final chunk if it has content
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunk_data = self._create_chunk_metadata(
                chunk_text, chunk_index, document_metadata,
                len(current_chunk), current_length
            )
            chunks.append(chunk_data)
        
        return chunks
    
    def _create_word_chunks(self, text: str, document_metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Fallback word-based chunking"""
        words = text.split()
        chunks = []
        chunk_index = 0
        
        i = 0
        while i < len(words):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            chunk_data = self._create_chunk_metadata(
                chunk_text, chunk_index, document_metadata,
                len(chunk_words), len(chunk_words)
            )
            chunks.append(chunk_data)
            
            # Move forward with overlap
            i += self.chunk_size - self.chunk_overlap
            chunk_index += 1
        
        return chunks
    
    def _calculate_sentence_overlap(self, sentences: List[str]) -> List[str]:
        """Calculate which sentences to overlap based on word count"""
        if not sentences:
            return []
        
        overlap_words = 0
        overlap_sentences = []
        
        # Start from the end and work backwards
        for sentence in reversed(sentences):
            sentence_words = len(sentence.split())
            if overlap_words + sentence_words <= self.chunk_overlap:
                overlap_sentences.insert(0, sentence)
                overlap_words += sentence_words
            else:
                break
        
        return overlap_sentences
    
    def _create_chunk_metadata(self, text: str, index: int, doc_metadata: Dict[str, Any], 
                             sentence_count: int, word_count: int) -> Dict[str, Any]:
        """Create comprehensive metadata for a chunk"""
        return {
            'text': text,
            'chunk_index': index,
            'word_count': word_count,
            'sentence_count': sentence_count,
            'char_count': len(text),
            'document_metadata': doc_metadata,
            'chunk_hash': hashlib.md5(text.encode('utf-8')).hexdigest(),
            'created_at': datetime.now(timezone.utc).isoformat(),
            'chunk_type': 'sentence' if sentence_count > 1 else 'word'
        }


class DocumentProcessor:
    """Process different document types with Windows-specific optimizations"""
    
    def __init__(self):
        self.logger = WindowsLogger()
        self.file_handler = WindowsFileHandler()
        self.supported_extensions = {'.txt', '.csv', '.json', '.pdf', '.docx'}
    
    def process_document(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """
        Process a single document and extract text content
        
        Args:
            file_path: Path to the document to process
            
        Returns:
            Dictionary with text content and metadata, or None if processing fails
        """
        try:
            # Check if file is locked
            if self.file_handler.is_file_locked(file_path):
                self.logger.warning(f"File is locked: {file_path}")
                return None
            
            # Get file metadata
            metadata = self.file_handler.get_file_metadata(file_path)
            
            if 'error' in metadata:
                self.logger.error(f"Error getting metadata for {file_path}: {metadata['error']}")
                return None
            
            # Check file extension
            extension = file_path.suffix.lower()
            if extension not in self.supported_extensions:
                self.logger.warning(f"Unsupported file type: {extension}")
                return None
            
            # Process based on file type
            content = None
            processing_metadata = {}
            
            try:
                if extension == '.txt':
                    content = self._process_text_file(file_path)
                elif extension == '.csv':
                    content, processing_metadata = self._process_csv_file(file_path)
                elif extension == '.json':
                    content, processing_metadata = self._process_json_file(file_path)
                elif extension == '.pdf':
                    content, processing_metadata = self._process_pdf_file(file_path)
                elif extension == '.docx':
                    content, processing_metadata = self._process_docx_file(file_path)
                
                if content:
                    return {
                        'content': content,
                        'metadata': {**metadata, **processing_metadata}
                    }
                else:
                    self.logger.warning(f"No content extracted from {file_path}")
                    return None
                    
            except Exception as e:
                self.logger.error(f"Error processing {file_path}: {str(e)}")
                self.logger.debug(traceback.format_exc())
                return None
                
        except Exception as e:
            self.logger.error(f"Unexpected error processing {file_path}: {str(e)}")
            return None
    
    def _process_text_file(self, file_path: Path) -> Optional[str]:
        """Process plain text files"""
        content = self.file_handler.safe_read_text_file(file_path)
        if content:
            # Basic text cleaning
            content = content.strip()
            # Remove excessive whitespace but preserve structure
            import re
            content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)  # Max 2 consecutive newlines
            content = re.sub(r'[ \t]+', ' ', content)  # Normalize spaces and tabs
        return content
    
    def _process_csv_file(self, file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        """Process CSV files"""
        try:
            # Try different encodings and delimiters
            encodings = ['utf-8', 'utf-8-sig', 'cp1252', 'iso-8859-1']
            delimiters = [',', ';', '\t', '|']
            
            df = None
            used_encoding = None
            used_delimiter = None
            
            for encoding in encodings:
                for delimiter in delimiters:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding, delimiter=delimiter, 
                                       on_bad_lines='skip', low_memory=False)
                        if len(df.columns) > 1:  # Good sign we found the right delimiter
                            used_encoding = encoding
                            used_delimiter = delimiter
                            break
                    except Exception:
                        continue
                if df is not None:
                    break
            
            if df is None or df.empty:
                return None, {}
            
            # Convert DataFrame to text
            content_parts = []
            
            # Add column headers
            content_parts.append("Columns: " + ", ".join(df.columns))
            
            # Add rows (limit to avoid memory issues)
            max_rows = min(1000, len(df))  # Limit for memory management
            for idx, row in df.head(max_rows).iterrows():
                row_text = " | ".join([f"{col}: {str(val)}" for col, val in row.items() if pd.notna(val)])
                if row_text:
                    content_parts.append(f"Row {idx + 1}: {row_text}")
            
            content = "\n".join(content_parts)
            
            metadata = {
                'csv_rows': len(df),
                'csv_columns': len(df.columns),
                'csv_encoding': used_encoding,
                'csv_delimiter': used_delimiter,
                'csv_column_names': list(df.columns),
                'processed_rows': max_rows
            }
            
            return content, metadata
            
        except Exception as e:
            self.logger.error(f"CSV processing error: {e}")
            return None, {}
    
    def _process_json_file(self, file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        """Process JSON files"""
        try:
            content_text = self.file_handler.safe_read_text_file(file_path)
            if not content_text:
                return None, {}
            
            json_data = json.loads(content_text)
            
            # Convert JSON to readable text
            def json_to_text(obj, prefix=""):
                text_parts = []
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            text_parts.append(f"{prefix}{key}:")
                            text_parts.extend(json_to_text(value, prefix + "  "))
                        else:
                            text_parts.append(f"{prefix}{key}: {value}")
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        if isinstance(item, (dict, list)):
                            text_parts.append(f"{prefix}Item {i + 1}:")
                            text_parts.extend(json_to_text(item, prefix + "  "))
                        else:
                            text_parts.append(f"{prefix}Item {i + 1}: {item}")
                else:
                    text_parts.append(f"{prefix}{obj}")
                return text_parts
            
            content_lines = json_to_text(json_data)
            content = "\n".join(content_lines)
            
            metadata = {
                'json_type': type(json_data).__name__,
                'json_size': len(content_text),
                'json_structure_depth': self._calculate_json_depth(json_data)
            }
            
            return content, metadata
            
        except Exception as e:
            self.logger.error(f"JSON processing error: {e}")
            return None, {}
    
    def _process_pdf_file(self, file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        """Process PDF files"""
        if not PDF_AVAILABLE:
            self.logger.error("PyPDF2 not available for PDF processing")
            return None, {}
        
        try:
            content_parts = []
            metadata = {'pdf_pages': 0, 'pdf_errors': []}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['pdf_pages'] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            content_parts.append(f"Page {page_num + 1}:\n{text}")
                    except Exception as e:
                        error_msg = f"Error extracting page {page_num + 1}: {str(e)}"
                        metadata['pdf_errors'].append(error_msg)
                        self.logger.warning(error_msg)
            
            content = "\n\n".join(content_parts) if content_parts else None
            metadata['pdf_extracted_pages'] = len(content_parts)
            
            return content, metadata
            
        except Exception as e:
            self.logger.error(f"PDF processing error: {e}")
            return None, {}
    
    def _process_docx_file(self, file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        """Process DOCX files"""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available for DOCX processing")
            return None, {}
        
        try:
            doc = DocxDocument(file_path)
            content_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_parts.append(text)
            
            # Extract tables
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    if row_data:
                        table_content.append(" | ".join(row_data))
            
            if table_content:
                content_parts.extend(["Tables:", *table_content])
            
            content = "\n".join(content_parts) if content_parts else None
            
            metadata = {
                'docx_paragraphs': len(doc.paragraphs),
                'docx_tables': len(doc.tables),
                'docx_has_images': len(doc.inline_shapes) > 0
            }
            
            return content, metadata
            
        except Exception as e:
            self.logger.error(f"DOCX processing error: {e}")
            return None, {}
    
    def _calculate_json_depth(self, obj, depth=0):
        """Calculate maximum depth of JSON structure"""
        if isinstance(obj, dict):
            return max([self._calculate_json_depth(v, depth + 1) for v in obj.values()], default=depth)
        elif isinstance(obj, list):
            return max([self._calculate_json_depth(item, depth + 1) for item in obj], default=depth)
        else:
            return depth


class QdrantIndexer:
    """Qdrant vector database indexer with Windows optimizations"""
    
    def __init__(self, host: str = "localhost", port: int = 6333, 
                 collection_name: str = "documents"):
        self.host = host
        self.port = port
        self.collection_name = collection_name
        self.logger = WindowsLogger()
        
        # Initialize Qdrant client
        try:
            self.client = QdrantClient(host=host, port=port)
            self.logger.info(f"Connected to Qdrant at {host}:{port}")
        except Exception as e:
            self.logger.error(f"Failed to connect to Qdrant: {e}")
            raise
        
        # Initialize sentence transformer
        try:
            self.model = SentenceTransformer('all-MiniLM-L6-v2')
            self.vector_size = self.model.get_sentence_embedding_dimension()
            self.logger.info(f"Loaded embedding model with dimension: {self.vector_size}")
        except Exception as e:
            self.logger.error(f"Failed to load embedding model: {e}")
            raise
        
        # Ensure collection exists
        self._ensure_collection()
    
    def _ensure_collection(self):
        """Ensure the collection exists with proper configuration"""
        try:
            collections = self.client.get_collections().collections
            collection_exists = any(col.name == self.collection_name for col in collections)
            
            if not collection_exists:
                self.client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config=VectorParams(
                        size=self.vector_size,
                        distance=Distance.COSINE
                    )
                )
                self.logger.info(f"Created collection: {self.collection_name}")
            else:
                self.logger.info(f"Collection exists: {self.collection_name}")
        except Exception as e:
            self.logger.error(f"Error ensuring collection: {e}")
            raise
    
    def index_document_chunks(self, chunks: List[Dict[str, Any]], 
                            batch_size: int = 100) -> Dict[str, Any]:
        """
        Index document chunks into Qdrant with batch processing
        
        Args:
            chunks: List of chunk dictionaries from DocumentChunker
            batch_size: Number of chunks to process in each batch
            
        Returns:
            Dictionary with indexing results and statistics
        """
        if not chunks:
            return {'success': False, 'message': 'No chunks to index'}
        
        start_time = time.time()
        total_chunks = len(chunks)
        successful_chunks = 0
        failed_chunks = 0
        
        self.logger.info(f"Starting indexing of {total_chunks} chunks")
        
        try:
            # Process in batches to manage memory
            for batch_start in range(0, total_chunks, batch_size):
                batch_end = min(batch_start + batch_size, total_chunks)
                batch_chunks = chunks[batch_start:batch_end]
                
                batch_result = self._process_batch(batch_chunks, batch_start)
                successful_chunks += batch_result['successful']
                failed_chunks += batch_result['failed']
                
                # Log progress
                progress = ((batch_end) / total_chunks) * 100
                self.logger.info(f"Processed batch {batch_start}-{batch_end}: {progress:.1f}% complete")
                
                # Force garbage collection to manage memory
                if batch_start > 0 and batch_start % (batch_size * 5) == 0:
                    gc.collect()
        
        except Exception as e:
            self.logger.error(f"Batch processing error: {e}")
            return {
                'success': False,
                'message': str(e),
                'processed_chunks': successful_chunks,
                'failed_chunks': failed_chunks
            }
        
        end_time = time.time()
        duration = end_time - start_time
        
        result = {
            'success': True,
            'total_chunks': total_chunks,
            'successful_chunks': successful_chunks,
            'failed_chunks': failed_chunks,
            'processing_time': duration,
            'chunks_per_second': successful_chunks / duration if duration > 0 else 0,
            'collection_name': self.collection_name
        }
        
        self.logger.info(f"Indexing complete: {successful_chunks}/{total_chunks} chunks indexed in {duration:.2f}s")
        
        return result
    
    def _process_batch(self, batch_chunks: List[Dict[str, Any]], batch_offset: int) -> Dict[str, int]:
        """Process a batch of chunks"""
        successful = 0
        failed = 0
        
        try:
            # Extract texts and generate embeddings
            texts = [chunk['text'] for chunk in batch_chunks]
            embeddings = self.model.encode(texts, batch_size=32, show_progress_bar=False)
            
            # Create points for Qdrant
            points = []
            for i, (chunk, embedding) in enumerate(zip(batch_chunks, embeddings)):
                try:
                    point_id = batch_offset + i
                    
                    # Prepare payload with all metadata
                    payload = {
                        'text': chunk['text'],
                        'chunk_index': chunk['chunk_index'],
                        'word_count': chunk['word_count'],
                        'sentence_count': chunk.get('sentence_count', 1),
                        'char_count': chunk['char_count'],
                        'chunk_hash': chunk['chunk_hash'],
                        'chunk_type': chunk.get('chunk_type', 'text'),
                        'created_at': chunk['created_at'],
                        # Document metadata
                        'file_path': chunk['document_metadata']['file_path'],
                        'file_name': chunk['document_metadata']['file_name'],
                        'file_extension': chunk['document_metadata']['file_extension'],
                        'file_size': chunk['document_metadata']['file_size'],
                        'file_created_time': chunk['document_metadata']['created_time'],
                        'file_modified_time': chunk['document_metadata']['modified_time'],
                        'mime_type': chunk['document_metadata']['mime_type']
                    }
                    
                    # Add any additional processing metadata
                    for key, value in chunk['document_metadata'].items():
                        if key not in payload and not key.startswith('file_'):
                            payload[f'doc_{key}'] = value
                    
                    point = PointStruct(
                        id=point_id,
                        vector=embedding.tolist(),
                        payload=payload
                    )
                    points.append(point)
                    
                except Exception as e:
                    self.logger.error(f"Error creating point for chunk {i}: {e}")
                    failed += 1
            
            # Upload batch to Qdrant
            if points:
                self.client.upsert(
                    collection_name=self.collection_name,
                    points=points
                )
                successful += len(points)
                
        except Exception as e:
            self.logger.error(f"Batch processing error: {e}")
            failed += len(batch_chunks)
        
        return {'successful': successful, 'failed': failed}
    
    def search_similar(self, query_text: str, limit: int = 10, 
                      score_threshold: float = 0.7) -> List[Dict[str, Any]]:
        """Search for similar document chunks"""
        try:
            # Generate query embedding
            query_embedding = self.model.encode([query_text])[0]
            
            # Search in Qdrant
            search_result = self.client.search(
                collection_name=self.collection_name,
                query_vector=query_embedding.tolist(),
                limit=limit,
                score_threshold=score_threshold
            )
            
            results = []
            for hit in search_result:
                results.append({
                    'id': hit.id,
                    'score': hit.score,
                    'text': hit.payload.get('text', ''),
                    'file_path': hit.payload.get('file_path', ''),
                    'file_name': hit.payload.get('file_name', ''),
                    'chunk_index': hit.payload.get('chunk_index', 0),
                    'metadata': hit.payload
                })
            
            return results
            
        except Exception as e:
            self.logger.error(f"Search error: {e}")
            return []
    
    def get_collection_info(self) -> Dict[str, Any]:
        """Get information about the collection"""
        try:
            info = self.client.get_collection(self.collection_name)
            return {
                'name': info.config.collection_name,
                'points_count': info.points_count,
                'vectors_count': info.vectors_count,
                'indexed_vectors_count': info.indexed_vectors_count,
                'status': info.status,
                'optimizer_status': info.optimizer_status,
                'vector_size': info.config.params.vectors.size,
                'distance': info.config.params.vectors.distance
            }
        except Exception as e:
            self.logger.error(f"Error getting collection info: {e}")
            return {}


class DocumentIndexingSystem:
    """Complete document indexing system with Windows optimizations"""
    
    def __init__(self, 
                 qdrant_host: str = "localhost",
                 qdrant_port: int = 6333,
                 collection_name: str = "documents",
                 chunk_size: int = 800,
                 chunk_overlap: int = 100):
        
        self.logger = WindowsLogger(name="document_indexing_system")
        
        # Initialize components
        self.processor = DocumentProcessor()
        self.chunker = DocumentChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        self.indexer = QdrantIndexer(
            host=qdrant_host, 
            port=qdrant_port, 
            collection_name=collection_name
        )
        
        self.logger.info("Document indexing system initialized")
    
    def index_directory(self, directory_path: str, recursive: bool = True,
                       file_patterns: List[str] = None) -> Dict[str, Any]:
        """
        Index all supported documents in a directory
        
        Args:
            directory_path: Path to directory to index
            recursive: Whether to search subdirectories
            file_patterns: List of file patterns to match (e.g., ['*.txt', '*.pdf'])
            
        Returns:
            Dictionary with indexing results and statistics
        """
        directory = Path(directory_path)
        
        if not directory.exists():
            return {'success': False, 'message': f'Directory does not exist: {directory_path}'}
        
        self.logger.info(f"Starting directory indexing: {directory_path}")
        
        # Find files to process
        files_to_process = self._find_files_to_process(directory, recursive, file_patterns)
        
        if not files_to_process:
            return {'success': False, 'message': 'No supported files found to index'}
        
        self.logger.info(f"Found {len(files_to_process)} files to process")
        
        # Process files
        results = {
            'success': True,
            'directory': str(directory),
            'total_files': len(files_to_process),
            'processed_files': 0,
            'failed_files': 0,
            'total_chunks': 0,
            'indexed_chunks': 0,
            'failed_chunks': 0,
            'processing_time': 0,
            'file_results': []
        }
        
        start_time = time.time()
        
        for i, file_path in enumerate(files_to_process):
            file_result = self.index_file(file_path)
            results['file_results'].append(file_result)
            
            if file_result['success']:
                results['processed_files'] += 1
                results['total_chunks'] += file_result.get('total_chunks', 0)
                results['indexed_chunks'] += file_result.get('indexed_chunks', 0)
                results['failed_chunks'] += file_result.get('failed_chunks', 0)
            else:
                results['failed_files'] += 1
            
            # Log progress
            progress = ((i + 1) / len(files_to_process)) * 100
            self.logger.info(f"Progress: {progress:.1f}% ({i + 1}/{len(files_to_process)} files)")
        
        results['processing_time'] = time.time() - start_time
        
        self.logger.info(f"Directory indexing complete: {results['processed_files']}/{results['total_files']} files, "
                        f"{results['indexed_chunks']} chunks indexed in {results['processing_time']:.2f}s")
        
        return results
    
    def index_file(self, file_path: str) -> Dict[str, Any]:
        """
        Index a single file
        
        Args:
            file_path: Path to file to index
            
        Returns:
            Dictionary with file processing results
        """
        file_path = Path(file_path)
        
        self.logger.info(f"Processing file: {file_path}")
        
        start_time = time.time()
        
        try:
            # Process document
            document_result = self.processor.process_document(file_path)
            
            if not document_result:
                return {
                    'success': False,
                    'file_path': str(file_path),
                    'message': 'Failed to process document',
                    'processing_time': time.time() - start_time
                }
            
            # Chunk document
            chunks = self.chunker.chunk_text(
                document_result['content'],
                document_result['metadata']
            )
            
            if not chunks:
                return {
                    'success': False,
                    'file_path': str(file_path),
                    'message': 'No chunks generated from document',
                    'processing_time': time.time() - start_time
                }
            
            # Index chunks
            indexing_result = self.indexer.index_document_chunks(chunks)
            
            result = {
                'success': indexing_result['success'],
                'file_path': str(file_path),
                'total_chunks': len(chunks),
                'indexed_chunks': indexing_result.get('successful_chunks', 0),
                'failed_chunks': indexing_result.get('failed_chunks', 0),
                'processing_time': time.time() - start_time,
                'file_size': document_result['metadata'].get('file_size', 0),
                'content_length': len(document_result['content'])
            }
            
            self.logger.info(f"File processed: {file_path.name} - "
                           f"{result['indexed_chunks']}/{result['total_chunks']} chunks indexed")
            
            return result
            
        except Exception as e:
            self.logger.error(f"Error processing file {file_path}: {e}")
            return {
                'success': False,
                'file_path': str(file_path),
                'message': str(e),
                'processing_time': time.time() - start_time
            }
    
    def _find_files_to_process(self, directory: Path, recursive: bool, 
                             file_patterns: List[str]) -> List[Path]:
        """Find files to process based on patterns and supported extensions"""
        files = []
        supported_extensions = {'.txt', '.csv', '.json', '.pdf', '.docx'}
        
        try:
            if recursive:
                search_pattern = "**/*"
            else:
                search_pattern = "*"
            
            for file_path in directory.glob(search_pattern):
                if file_path.is_file():
                    # Check extension
                    if file_path.suffix.lower() in supported_extensions:
                        # Check file patterns if specified
                        if file_patterns:
                            if any(file_path.match(pattern) for pattern in file_patterns):
                                files.append(file_path)
                        else:
                            files.append(file_path)
                            
        except Exception as e:
            self.logger.error(f"Error finding files in {directory}: {e}")
        
        return files
    
    def search(self, query: str, limit: int = 10, score_threshold: float = 0.7) -> List[Dict[str, Any]]:
        """Search indexed documents"""
        return self.indexer.search_similar(query, limit, score_threshold)
    
    def get_system_status(self) -> Dict[str, Any]:
        """Get system status and statistics"""
        try:
            collection_info = self.indexer.get_collection_info()
            return {
                'status': 'operational',
                'collection': collection_info,
                'supported_extensions': list(self.processor.supported_extensions),
                'chunk_size': self.chunker.chunk_size,
                'chunk_overlap': self.chunker.chunk_overlap,
                'embedding_model': 'all-MiniLM-L6-v2',
                'vector_dimension': self.indexer.vector_size
            }
        except Exception as e:
            return {
                'status': 'error',
                'message': str(e)
            }